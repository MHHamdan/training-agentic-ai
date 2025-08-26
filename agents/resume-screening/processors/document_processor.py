import os
import logging
from typing import Optional, Dict, Any
import asyncio
from pathlib import Path
import PyPDF2
import docx
import pytesseract
from PIL import Image
import io
from langsmith import traceable

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.txt', '.doc', '.rtf'}
        self.max_file_size = 10 * 1024 * 1024  # 10MB
    
    @traceable(name="extract_resume_text", metadata={"component": "document_processor"})
    async def extract_text(self, file_path: str) -> Optional[str]:
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                logger.error(f"File not found: {file_path}")
                return None
            
            if file_path.stat().st_size > self.max_file_size:
                logger.error(f"File too large: {file_path}")
                return None
            
            file_extension = file_path.suffix.lower()
            
            if file_extension == '.pdf':
                return await self._extract_pdf_text(file_path)
            elif file_extension in ['.docx', '.doc']:
                return await self._extract_docx_text(file_path)
            elif file_extension in ['.txt', '.rtf']:
                return await self._extract_text_file(file_path)
            else:
                logger.error(f"Unsupported file format: {file_extension}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            return None
    
    async def _extract_pdf_text(self, file_path: Path) -> Optional[str]:
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if not text.strip():
                logger.info("No text found in PDF, attempting OCR...")
                text = await self._perform_ocr_on_pdf(file_path)
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            return None
    
    async def _extract_docx_text(self, file_path: Path) -> Optional[str]:
        try:
            doc = docx.Document(str(file_path))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            return None
    
    async def _extract_text_file(self, file_path: Path) -> Optional[str]:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read().strip()
        except Exception as e:
            logger.error(f"Error reading text file: {str(e)}")
            return None
    
    async def _perform_ocr_on_pdf(self, file_path: Path) -> str:
        try:
            import pdf2image
            
            images = pdf2image.convert_from_path(str(file_path))
            text = ""
            
            for i, image in enumerate(images):
                try:
                    page_text = pytesseract.image_to_string(image)
                    text += f"\n--- Page {i+1} ---\n{page_text}"
                except Exception as e:
                    logger.error(f"OCR error on page {i+1}: {str(e)}")
            
            return text.strip()
        except Exception as e:
            logger.error(f"OCR processing error: {str(e)}")
            return ""
    
    @traceable(name="extract_metadata", metadata={"component": "document_processor"})
    async def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        try:
            file_path = Path(file_path)
            stat = file_path.stat()
            
            metadata = {
                "filename": file_path.name,
                "file_size": stat.st_size,
                "file_size_mb": round(stat.st_size / (1024 * 1024), 2),
                "file_extension": file_path.suffix.lower(),
                "created_time": stat.st_ctime,
                "modified_time": stat.st_mtime,
            }
            
            if file_path.suffix.lower() == '.pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata["pages"] = len(pdf_reader.pages)
                    if pdf_reader.metadata:
                        metadata["pdf_metadata"] = {
                            "title": pdf_reader.metadata.get('/Title', ''),
                            "author": pdf_reader.metadata.get('/Author', ''),
                            "subject": pdf_reader.metadata.get('/Subject', ''),
                        }
            
            return metadata
        except Exception as e:
            logger.error(f"Error extracting metadata: {str(e)}")
            return {}
    
    @traceable(name="sanitize_text", metadata={"component": "document_processor"})
    async def sanitize_text(self, text: str) -> str:
        import re
        
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        text = re.sub(email_pattern, '[EMAIL]', text)
        
        phone_pattern = r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b'
        text = re.sub(phone_pattern, '[PHONE]', text)
        
        ssn_pattern = r'\b\d{3}-\d{2}-\d{4}\b'
        text = re.sub(ssn_pattern, '[SSN]', text)
        
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        
        return text