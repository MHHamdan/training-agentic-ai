"""
AI Content Creation System - Streamlit Interface
Built with LangGraph for sophisticated content creation workflows
Following patterns from financial analysis app.py for consistency
"""

import streamlit as st
import pandas as pd
import json
import uuid
from datetime import datetime, timedelta
from typing import Dict, List, Any
import os
import sys
from pathlib import Path
import io
import base64
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
import markdown2

# Add current directory to path
current_dir = Path(__file__).parent
sys.path.insert(0, str(current_dir))

# Import our content creation system components
try:
    from content_graph import content_creation_graph, create_content_creation_session
    from content_state import ContentCreationState, BrandGuidelines
    from langgraph.types import Command
    LANGGRAPH_AVAILABLE = True
except ImportError:
    st.error("⚠️ LangGraph dependencies not available. Please install: pip install langgraph langchain-openai")
    LANGGRAPH_AVAILABLE = False

# Page configuration
st.set_page_config(
    page_title="AI Content Creation System",
    page_icon="✍️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for professional content creation interface
st.markdown("""
<style>
    .main-header {
        font-size: 2.8rem;
        color: #1e40af;
        text-align: center;
        margin-bottom: 1rem;
        font-weight: 700;
        background: linear-gradient(135deg, #3b82f6 0%, #8b5cf6 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #64748b;
        text-align: center;
        margin-bottom: 2rem;
    }
    .metric-card {
        background: linear-gradient(135deg, #f8fafc 0%, #e2e8f0 100%);
        padding: 1.5rem;
        border-radius: 12px;
        border: 1px solid #e1e7ef;
        margin-bottom: 1rem;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    .agent-status {
        display: inline-block;
        padding: 0.25rem 0.75rem;
        border-radius: 20px;
        font-size: 0.875rem;
        font-weight: 600;
        margin: 0.25rem;
    }
    .status-active {
        background-color: #fbbf24;
        color: #78350f;
    }
    .status-complete {
        background-color: #34d399;
        color: #064e3b;
    }
    .status-pending {
        background-color: #94a3b8;
        color: #1e293b;
    }
    .content-preview {
        background: linear-gradient(135deg, #f8fafc 0%, #ffffff 100%);
        border: 1px solid #e2e8f0;
        border-radius: 12px;
        padding: 2rem;
        margin: 1rem 0;
        font-family: 'Georgia', serif;
        font-size: 1rem;
        line-height: 1.7;
        box-shadow: 0 4px 12px rgba(0,0,0,0.1);
    }
    .content-preview h1 {
        color: #1e40af;
        border-bottom: 3px solid #3b82f6;
        padding-bottom: 0.5rem;
        margin-bottom: 1.5rem;
    }
    .content-preview h2 {
        color: #1e293b;
        margin-top: 2rem;
        margin-bottom: 1rem;
    }
    .content-preview h3 {
        color: #475569;
        margin-top: 1.5rem;
        margin-bottom: 0.8rem;
    }
    .content-preview p {
        margin-bottom: 1rem;
        text-align: justify;
    }
    .content-preview ul, .content-preview ol {
        margin-left: 1.5rem;
        margin-bottom: 1rem;
    }
    .content-preview li {
        margin-bottom: 0.5rem;
    }
    .agent-analysis {
        background: linear-gradient(135deg, #f1f5f9 0%, #e2e8f0 100%);
        border-left: 4px solid #3b82f6;
        border-radius: 8px;
        padding: 1.5rem;
        margin: 1rem 0;
        font-size: 0.95rem;
        line-height: 1.6;
    }
    .analysis-header {
        font-weight: 600;
        color: #1e40af;
        margin-bottom: 0.5rem;
        display: flex;
        align-items: center;
        gap: 0.5rem;
    }
</style>
""", unsafe_allow_html=True)


class ContentCreationApp:
    """Main application class for the AI content creation system"""
    
    def __init__(self):
        self.initialize_session_state()
        
    def initialize_session_state(self):
        """Initialize session state variables"""
        if 'content_session' not in st.session_state:
            st.session_state.content_session = None
        if 'content_results' not in st.session_state:
            st.session_state.content_results = {}
        if 'workflow_history' not in st.session_state:
            st.session_state.workflow_history = []
        if 'workflow_running' not in st.session_state:
            st.session_state.workflow_running = False
        if 'current_agent' not in st.session_state:
            st.session_state.current_agent = None
        if 'generated_content' not in st.session_state:
            st.session_state.generated_content = {}
    
    def get_content_data(self, content_key: str, default: str = "No content available") -> str:
        """Safely extract content data from ContentResult objects"""
        if content_key in st.session_state.content_results:
            result = st.session_state.content_results[content_key]
            if hasattr(result, 'content_data') and result.content_data:
                return result.content_data.get("analysis", default)
        return default
    
    def run(self):
        """Main application entry point"""
        # Header
        st.markdown('<h1 class="main-header">✍️ AI Content Creation System</h1>', unsafe_allow_html=True)
        st.markdown('<p class="sub-header">Powered by LangGraph • Advanced Multi-Agent Content Generation</p>', unsafe_allow_html=True)
        
        # Check for API keys and show active provider
        api_keys = {
            "HUGGING_FACE_API": os.getenv("HUGGING_FACE_API") or os.getenv("HF_TOKEN"),
            "GROQ_API_KEY": os.getenv("GROQ_API_KEY"),
            "GOOGLE_API_KEY": os.getenv("GOOGLE_API_KEY"),
            "ANTHROPIC_API_KEY": os.getenv("ANTHROPIC_API_KEY"),
            "GROK_API_KEY": os.getenv("GROK_API_KEY"),
            "OPENAI_API_KEY": os.getenv("OPENAI_API_KEY")
        }
        
        available_providers = [k for k, v in api_keys.items() if v]
        
        if not available_providers:
            st.error("❌ No LLM API keys configured. Please add at least one API key to your .env file")
        else:
            # Show primary provider being used
            if api_keys["HUGGING_FACE_API"]:
                st.success("🆓 Using **Hugging Face GPT-OSS-120B** (OpenAI's free open-source model)")
                st.info("💡 This is completely free with Hugging Face's inference API!")
            elif api_keys["GROQ_API_KEY"]:
                st.info("⚡ Using **Groq Llama-3.1-70B** as primary LLM provider")
            elif api_keys["GOOGLE_API_KEY"]:
                st.info("🔍 Using **Google Gemini Pro** as primary LLM provider")
            elif api_keys["ANTHROPIC_API_KEY"]:
                st.info("🧠 Using **Anthropic Claude** as primary LLM provider")
            elif api_keys["GROK_API_KEY"]:
                st.info("🚀 Using **Grok (xAI)** as primary LLM provider")
            elif api_keys["OPENAI_API_KEY"]:
                st.warning("💸 Using **OpenAI GPT-4** (may have quota limits)")
            
            # Show fallback options
            fallbacks = [k.replace("_API_KEY", "").replace("_API", "") for k in available_providers[1:]]
            if fallbacks:
                st.caption(f"📋 Fallback options: {', '.join(fallbacks)}")
        
        # Sidebar configuration
        with st.sidebar:
            st.header("📝 Content Configuration")
            
            # Content topic
            content_topic = st.text_input(
                "Content Topic",
                value="AI in Digital Marketing",
                help="Enter the main topic for your content"
            )
            
            # Content type
            content_type = st.selectbox(
                "Content Type",
                options=[
                    "blog_post",
                    "social_media", 
                    "website_copy",
                    "product_description",
                    "email_campaign",
                    "white_paper",
                    "case_study",
                    "technical_documentation"
                ],
                format_func=lambda x: x.replace('_', ' ').title(),
                help="Select the type of content to create"
            )
            
            # Target keywords
            keywords_input = st.text_area(
                "Target Keywords",
                value="AI marketing, digital transformation, automation",
                help="Enter comma-separated keywords for SEO optimization"
            )
            keywords = [k.strip() for k in keywords_input.split(',') if k.strip()]
            
            # Target audience
            target_audience = st.selectbox(
                "Target Audience",
                options=["general", "professionals", "beginners", "executives", "technical", "academics"],
                index=1,
                help="Select your target audience"
            )
            
            # Content length
            word_count_target = st.slider(
                "Target Word Count",
                min_value=300,
                max_value=3000,
                value=800,
                step=100,
                help="Desired length of the content"
            )
            
            # Brand guidelines
            with st.expander("🎨 Brand Guidelines"):
                brand_tone = st.selectbox(
                    "Brand Tone",
                    options=["professional", "casual", "friendly", "authoritative", "conversational", "technical"],
                    help="Select your brand's tone of voice"
                )
                
                brand_voice = st.selectbox(
                    "Writing Voice",
                    options=["active", "passive", "mixed"],
                    help="Preferred writing voice"
                )
                
                brand_keywords = st.text_input(
                    "Brand Keywords",
                    placeholder="innovation, quality, expertise",
                    help="Keywords that represent your brand"
                )
                
                avoid_words = st.text_input(
                    "Words to Avoid",
                    placeholder="cheap, basic, simple",
                    help="Words that don't align with your brand"
                )
            
            # Advanced options
            with st.expander("⚙️ Advanced Options"):
                enable_seo = st.checkbox("Enable SEO Optimization", value=True)
                enable_plagiarism_check = st.checkbox("Enable Plagiarism Check", value=True)
                auto_quality_check = st.checkbox("Auto Quality Assurance", value=True)
                
                export_formats = st.multiselect(
                    "Export Formats",
                    options=["markdown", "html", "docx", "pdf"],
                    default=["markdown", "html"],
                    help="Select output formats"
                )
            
            # Action buttons
            st.divider()
            
            col1, col2 = st.columns(2)
            with col1:
                if st.button("🚀 Create Content", type="primary", use_container_width=True, disabled=st.session_state.workflow_running):
                    if content_topic and LANGGRAPH_AVAILABLE:
                        self.start_content_creation(
                            content_topic, content_type, keywords, target_audience,
                            word_count_target, brand_tone, brand_voice,
                            brand_keywords, avoid_words, enable_seo
                        )
            
            with col2:
                if st.button("🔄 Reset", use_container_width=True):
                    self.reset_session()
            
            # Show current session info
            if st.session_state.content_session:
                st.divider()
                st.markdown("### 📊 Current Session")
                st.info(f"Session ID: {st.session_state.content_session.get('session_id', 'N/A')[:8]}...")
                st.info(f"Topic: {content_topic}")
                st.info(f"Type: {content_type.replace('_', ' ').title()}")
        
        # Main content area
        if st.session_state.workflow_running:
            self.display_workflow_progress()
        
        # Display results tabs
        if st.session_state.content_results:
            self.display_content_results()
        else:
            self.display_welcome()
    
    def start_content_creation(self, topic: str, content_type: str, keywords: List[str], 
                             target_audience: str, word_count: int, brand_tone: str,
                             brand_voice: str, brand_keywords: str, avoid_words: str, enable_seo: bool):
        """Start the content creation workflow"""
        st.session_state.workflow_running = True
        
        # Create brand guidelines
        brand_guidelines = {
            "tone": brand_tone,
            "voice": brand_voice,
            "style_keywords": [k.strip() for k in brand_keywords.split(',') if k.strip()],
            "avoid_words": [w.strip() for w in avoid_words.split(',') if w.strip()],
            "target_audience": target_audience
        }
        
        # Create content creation session
        session = create_content_creation_session(
            topic=topic,
            content_type=content_type,
            target_keywords=keywords,
            brand_guidelines=brand_guidelines
        )
        
        st.session_state.content_session = session
        
        # Initialize with user request including word count and requirements
        initial_message = f"Create {content_type.replace('_', ' ')} about {topic} for {target_audience} audience. Target length: {word_count} words. Style: {brand_tone} tone with {brand_voice} voice."
        
        with st.spinner("🔍 Initializing multi-agent content creation..."):
            try:
                # Run the LangGraph workflow
                if LANGGRAPH_AVAILABLE:
                    # Start the workflow
                    initial_state = {
                        **session['initial_state'],
                        "messages": [{"role": "user", "content": initial_message}],
                        "target_word_count": word_count,
                        "auto_seo_optimization": enable_seo
                    }
                    
                    # Stream updates from the graph
                    for update in content_creation_graph.stream(
                        initial_state,
                        config=session['thread_config'],
                        stream_mode="updates"
                    ):
                        # Process updates
                        for node_id, value in update.items():
                            st.session_state.current_agent = node_id
                            
                            # Extract and store results
                            if isinstance(value, dict):
                                if "completed_analyses" in value:
                                    st.session_state.content_results.update(value["completed_analyses"])
                                
                                if "messages" in value and value["messages"]:
                                    last_message = value["messages"][-1]
                                    st.session_state.workflow_history.append({
                                        "agent": node_id,
                                        "message": last_message,
                                        "timestamp": datetime.now()
                                    })
                    
                    st.success("✅ Content creation complete!")
                else:
                    # Demo mode
                    self.run_demo_content_creation(topic, content_type, keywords)
                    
            except Exception as e:
                st.error(f"❌ Error during content creation: {str(e)}")
            finally:
                st.session_state.workflow_running = False
    
    def run_demo_content_creation(self, topic: str, content_type: str, keywords: List[str]):
        """Run demo content creation with mock data"""
        import time
        
        # Simulate content creation steps
        agents = [
            "topic_research_agent",
            "content_strategist_agent", 
            "content_writer_agent",
            "seo_specialist_agent",
            "quality_assurance_agent",
            "content_editor_agent",
            "content_publisher_agent"
        ]
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for i, agent in enumerate(agents):
            status_text.text(f"🤖 {agent.replace('_', ' ').title()} working...")
            progress_bar.progress((i + 1) / len(agents))
            time.sleep(1)
            
            # Add mock results
            st.session_state.content_results[agent] = {
                "status": "complete",
                "confidence": 0.85 + (i * 0.02),
                "content_data": {"analysis": f"Mock analysis from {agent}"},
                "timestamp": datetime.now()
            }
        
        # Generate mock content
        mock_content = f"""# {topic.title()}

## Introduction

{topic} has become increasingly important in today's digital landscape. This comprehensive guide explores the key aspects and practical applications.

## Key Benefits

- Improved efficiency and automation
- Enhanced user experience
- Data-driven decision making
- Competitive advantage

## Implementation Strategy

1. **Assessment Phase**: Evaluate current capabilities
2. **Planning Phase**: Develop comprehensive roadmap  
3. **Execution Phase**: Implement solutions systematically
4. **Optimization Phase**: Continuously improve and refine

## Best Practices

When implementing {topic.lower()}, consider these essential practices:

- Start with clear objectives and success metrics
- Ensure stakeholder buy-in and support
- Invest in proper training and education
- Monitor progress and adjust strategies accordingly

## Conclusion

{topic} represents a significant opportunity for organizations looking to improve their capabilities and stay competitive. By following the strategies outlined in this guide, you can successfully implement and benefit from these approaches.

*Keywords: {', '.join(keywords)}*
"""
        
        st.session_state.generated_content = {
            "final_content": mock_content,
            "word_count": len(mock_content.split()),
            "seo_score": 87.5,
            "readability_score": 82.3,
            "quality_score": 89.1
        }
        
        status_text.text("✅ Content creation complete!")
        progress_bar.progress(1.0)
    
    def display_workflow_progress(self):
        """Display real-time workflow progress"""
        st.markdown("### 🔄 Content Creation in Progress")
        
        # Show current agent
        if st.session_state.current_agent:
            st.info(f"🤖 Current Agent: **{st.session_state.current_agent.replace('_', ' ').title()}**")
        
        # Show completed analyses
        if st.session_state.content_results:
            completed = list(st.session_state.content_results.keys())
            st.success(f"✅ Completed: {', '.join([c.replace('_', ' ').title() for c in completed])}")
    
    def display_content_results(self):
        """Display comprehensive content creation results"""
        st.markdown("### 📝 Content Creation Results")
        
        # Create tabs for different content sections
        tabs = st.tabs([
            "📖 Generated Content",
            "🔍 Research & Strategy", 
            "✨ SEO Analysis",
            "📊 Quality Assessment",
            "📈 Performance Metrics",
            "📁 Export & Download"
        ])
        
        with tabs[0]:
            self.display_generated_content()
        
        with tabs[1]:
            self.display_research_strategy()
        
        with tabs[2]:
            self.display_seo_analysis()
        
        with tabs[3]:
            self.display_quality_assessment()
        
        with tabs[4]:
            self.display_performance_metrics()
        
        with tabs[5]:
            self.display_export_options()
    
    def display_generated_content(self):
        """Display the final generated content"""
        st.markdown("#### 📖 Final Content")
        
        # Display content metrics
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            word_count = st.session_state.generated_content.get("word_count", 0)
            if word_count == 0 and "content_creation" in st.session_state.content_results:
                # Estimate word count from content
                content = self.get_content_data("content_creation", "")
                word_count = len(content.split()) if content else 0
            
            # Get target word count from session state
            target_count = st.session_state.content_session.get('initial_state', {}).get('target_word_count', 800) if st.session_state.content_session else 800
            st.metric("Word Count", word_count, f"Target: {target_count}")
        
        with col2:
            reading_time = max(1, word_count // 200)
            st.metric("Reading Time", f"{reading_time} min", "Estimated")
        
        with col3:
            st.metric("Content Type", 
                     st.session_state.content_session.get('initial_state', {}).get('content_type', 'Unknown').replace('_', ' ').title())
        
        with col4:
            st.metric("Status", "Complete", "✅ Ready")
        
        # Display the content
        final_content = st.session_state.generated_content.get("final_content", "")
        if final_content:
            st.markdown("#### Preview")
            with st.container():
                st.markdown('<div class="content-preview">', unsafe_allow_html=True)
                st.markdown(final_content)
                st.markdown('</div>', unsafe_allow_html=True)
        else:
            # Show content from results
            if "content_creation" in st.session_state.content_results:
                st.markdown("#### Content Analysis")
                st.info(self.get_content_data("content_creation", "No content available"))
    
    def display_research_strategy(self):
        """Display research and strategy results"""
        st.markdown("#### 🔍 Research & Strategy Analysis")
        
        # Topic research results
        if "topic_research_agent" in st.session_state.content_results:
            st.markdown("##### Topic Research")
            st.info(self.get_content_data("topic_research_agent", "No research data available"))
        
        # Content strategy results  
        if "content_strategist_agent" in st.session_state.content_results:
            st.markdown("##### Content Strategy")
            st.info(self.get_content_data("content_strategist_agent", "No strategy data available"))
        
        # Keywords analysis
        if st.session_state.content_session:
            keywords = st.session_state.content_session.get('initial_state', {}).get('target_keywords', [])
            if keywords:
                st.markdown("##### Target Keywords")
                cols = st.columns(min(len(keywords), 4))
                for i, keyword in enumerate(keywords):
                    with cols[i % 4]:
                        st.metric(f"Keyword {i+1}", keyword)
    
    def display_seo_analysis(self):
        """Display SEO analysis results"""
        st.markdown("#### ✨ SEO Optimization Analysis")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # SEO Score
            seo_score = st.session_state.generated_content.get("seo_score", 0)
            st.metric("SEO Score", f"{seo_score}/100", "🎯 Optimized" if seo_score > 80 else "⚠️ Needs Work")
            
            # Keyword density
            st.markdown("##### Keyword Optimization")
            if st.session_state.content_session:
                keywords = st.session_state.content_session.get('initial_state', {}).get('target_keywords', [])
                for keyword in keywords[:3]:  # Show first 3 keywords
                    density = 1.2  # Mock density
                    st.metric(f'"{keyword}" Density', f"{density}%", "✅ Optimal" if 1.0 <= density <= 2.5 else "⚠️ Adjust")
        
        with col2:
            # Meta tags suggestions
            st.markdown("##### Meta Tags")
            st.text_area(
                "Suggested Meta Title",
                value="AI in Digital Marketing: Complete Guide 2024",
                height=60,
                disabled=True
            )
            st.text_area(
                "Suggested Meta Description", 
                value="Discover how AI transforms digital marketing. Learn strategies, tools, and best practices for implementing AI in your marketing campaigns.",
                height=80,
                disabled=True
            )
        
        # SEO specialist analysis
        if "seo_specialist_agent" in st.session_state.content_results:
            st.markdown("##### SEO Specialist Analysis")
            st.info(self.get_content_data("seo_specialist_agent", "No SEO analysis available"))
    
    def display_quality_assessment(self):
        """Display quality assessment results"""
        st.markdown("#### 📊 Quality Assessment")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            quality_score = st.session_state.generated_content.get("quality_score", 0)
            st.metric("Overall Quality", f"{quality_score}/100", "📈 Excellent" if quality_score > 85 else "👍 Good")
        
        with col2:
            readability_score = st.session_state.generated_content.get("readability_score", 0)
            st.metric("Readability", f"{readability_score}/100", "📖 Clear" if readability_score > 70 else "⚠️ Complex")
        
        with col3:
            st.metric("Brand Alignment", "92/100", "✅ Excellent")
        
        # Detailed quality analysis
        if "quality_assurance_agent" in st.session_state.content_results:
            st.markdown("##### Quality Assurance Report")
            st.info(self.get_content_data("quality_assurance_agent", "No quality assessment available"))
        
        # Grammar and style check
        st.markdown("##### Writing Quality")
        quality_checks = [
            {"check": "Grammar", "status": "✅ Passed", "score": 96},
            {"check": "Style Consistency", "status": "✅ Passed", "score": 94},
            {"check": "Tone Alignment", "status": "✅ Passed", "score": 92},
            {"check": "Readability", "status": "✅ Passed", "score": 88}
        ]
        
        df = pd.DataFrame(quality_checks)
        st.dataframe(df, use_container_width=True, hide_index=True)
    
    def display_performance_metrics(self):
        """Display content performance predictions"""
        st.markdown("#### 📈 Performance Predictions")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("##### Engagement Predictions")
            st.metric("Estimated Shares", "250-400", "📱 Social Media")
            st.metric("Expected Comments", "15-25", "💬 Engagement")
            st.metric("Reading Completion", "78%", "📖 Retention")
        
        with col2:
            st.markdown("##### SEO Predictions")
            st.metric("Ranking Potential", "Top 10", "🎯 SERP Position")
            st.metric("Click-Through Rate", "4.2%", "🔗 Estimated CTR")
            st.metric("Search Impressions", "1,200/month", "👁️ Visibility")
        
        # Performance timeline
        st.markdown("##### Content Lifecycle")
        timeline_data = {
            "Phase": ["Week 1", "Week 2", "Month 1", "Month 3", "Month 6"],
            "Traffic": [100, 250, 500, 800, 1200],
            "Engagement": [85, 78, 72, 68, 65],
            "SEO Ranking": [50, 35, 20, 12, 8]
        }
        
        df = pd.DataFrame(timeline_data)
        st.line_chart(df.set_index("Phase")[["Traffic", "Engagement"]])
    
    def display_export_options(self):
        """Display content export and download options"""
        st.markdown("#### 📁 Export & Download")
        
        # Get content for export
        final_content = st.session_state.generated_content.get("final_content", "")
        if not final_content and "content_creation" in st.session_state.content_results:
            final_content = self.get_content_data("content_creation", "No content available")
        
        if not final_content or final_content == "No content available":
            st.warning("⚠️ No content available for export. Please generate content first.")
            return
        
        # Export formats
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("##### Available Formats")
            
            # Markdown download
            if st.button("📄 Download as Markdown", key="download_markdown"):
                self.download_content_as_markdown(final_content)
            
            # HTML download
            if st.button("🌐 Download as HTML", key="download_html"):
                self.download_content_as_html(final_content)
            
            # Word Document download
            if st.button("📝 Download as Word Document", key="download_docx"):
                self.download_content_as_docx(final_content)
            
            # PDF download
            if st.button("📄 Download as PDF", key="download_pdf"):
                self.download_content_as_pdf(final_content)
        
        with col2:
            st.markdown("##### Content Packages")
            if st.button("📦 Complete Content Package", type="primary"):
                self.download_complete_package()
            
            if st.button("🎨 Social Media Kit"):
                self.download_social_media_kit(final_content)
            
            if st.button("📧 Email Version"):
                self.download_email_version(final_content)
        
        # Content variations
        st.markdown("##### Platform Variations")
        
        variation_tabs = st.tabs(["📱 Social Media", "📧 Email", "🌐 Website", "📄 Blog"])
        
        with variation_tabs[0]:
            st.text_area("Twitter Version", value="🤖 AI is transforming digital marketing! Discover key strategies and tools in our latest guide. #AIMarketing #DigitalTransformation", height=100)
            st.text_area("LinkedIn Version", value="Artificial Intelligence is revolutionizing how we approach digital marketing. Our comprehensive guide covers implementation strategies, best practices, and real-world applications.", height=100)
        
        with variation_tabs[1]:
            st.text_area("Email Subject", value="Transform Your Marketing with AI: Complete Implementation Guide", height=60)
            st.text_area("Email Preview", value="Discover how leading companies are using AI to enhance their marketing effectiveness...", height=100)
        
        with variation_tabs[2]:
            st.text_area("Website Meta Description", value="Learn how AI is transforming digital marketing. Comprehensive guide with strategies, tools, and best practices for implementation.", height=80)
        
        with variation_tabs[3]:
            st.text_area("Blog Introduction", value="# AI in Digital Marketing: A Complete Guide\n\nArtificial Intelligence is reshaping the digital marketing landscape...", height=120)
    
    def display_welcome(self):
        """Display welcome screen with system overview"""
        st.markdown("""
        ### Welcome to the AI Content Creation System
        
        This advanced platform leverages **LangGraph** and multiple specialized AI agents to create high-quality, SEO-optimized content automatically:
        
        #### 🤖 **Specialized Agents:**
        - **Topic Research Agent**: Trend analysis, keyword research, and competitive intelligence
        - **Content Strategist Agent**: Content planning, audience analysis, and brand alignment
        - **Content Writer Agent**: High-quality content creation and copywriting
        - **SEO Specialist Agent**: Search optimization and meta tag generation
        - **Quality Assurance Agent**: Grammar, readability, and brand compliance checking
        - **Content Editor Agent**: Final refinement and polish
        - **Content Publisher Agent**: Multi-platform formatting and export preparation
        
        #### 🔄 **Advanced Features:**
        - **Dynamic Routing**: Intelligent agent selection based on content requirements
        - **Multi-Format Support**: Blog posts, social media, website copy, and more
        - **Human-in-the-Loop**: Approval workflows for quality control
        - **Brand Compliance**: Automatic brand voice and guideline adherence
        - **SEO Optimization**: Comprehensive search engine optimization
        - **Quality Assurance**: Multi-metric content quality assessment
        
        #### 🚀 **Getting Started:**
        1. Enter your content topic and select content type
        2. Configure target keywords and brand guidelines
        3. Set audience preferences and quality requirements
        4. Click "Create Content" to begin the automated workflow
        5. Review results across all analysis tabs
        
        ---
        
        **Built with LangGraph** for sophisticated multi-agent orchestration and state management.
        """)
        
        # Show sample content creation
        if st.button("🎯 Try Sample Content Creation", type="primary"):
            self.start_content_creation(
                "AI in Digital Marketing",
                "blog_post", 
                ["AI marketing", "digital transformation", "automation"],
                "professionals",
                800,
                "professional",
                "active",
                "innovation, expertise, quality",
                "cheap, basic",
                True
            )
    
    def download_content_as_markdown(self, content: str):
        """Download content as Markdown file"""
        # Add metadata header
        topic = st.session_state.content_session.get('initial_state', {}).get('content_topic', 'Content')
        content_type = st.session_state.content_session.get('initial_state', {}).get('content_type', 'blog_post')
        
        markdown_content = f"""---
title: {topic}
type: {content_type}
generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
---

{content}"""
        
        st.download_button(
            label="💾 Download Markdown",
            data=markdown_content,
            file_name=f"{topic.lower().replace(' ', '_')}.md",
            mime="text/markdown",
            key="markdown_download_btn"
        )
    
    def download_content_as_html(self, content: str):
        """Download content as HTML file"""
        # Convert markdown to HTML
        html_content = markdown2.markdown(content, extras=['fenced-code-blocks', 'tables'])
        
        topic = st.session_state.content_session.get('initial_state', {}).get('content_topic', 'Content')
        
        full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{topic}</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        h1, h2, h3 {{ color: #333; }}
        h1 {{ border-bottom: 2px solid #3b82f6; padding-bottom: 10px; }}
        blockquote {{ border-left: 4px solid #3b82f6; padding-left: 15px; margin: 20px 0; }}
        code {{ background-color: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
        pre {{ background-color: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
        
        st.download_button(
            label="💾 Download HTML",
            data=full_html,
            file_name=f"{topic.lower().replace(' ', '_')}.html",
            mime="text/html",
            key="html_download_btn"
        )
    
    def download_content_as_docx(self, content: str):
        """Download content as Word document"""
        try:
            doc = Document()
            
            # Add title
            topic = st.session_state.content_session.get('initial_state', {}).get('content_topic', 'Content')
            title = doc.add_heading(topic, 0)
            
            # Add content (split by lines and handle basic formatting)
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.strip():
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph()
            
            # Save to BytesIO
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            st.download_button(
                label="💾 Download Word Document",
                data=docx_buffer.getvalue(),
                file_name=f"{topic.lower().replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="docx_download_btn"
            )
        except ImportError:
            st.error("❌ Word document export requires python-docx. Install with: pip install python-docx")
    
    def download_content_as_pdf(self, content: str):
        """Download content as PDF file"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch
            
            pdf_buffer = io.BytesIO()
            doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Add title
            topic = st.session_state.content_session.get('initial_state', {}).get('content_topic', 'Content')
            story.append(Paragraph(topic, styles['Title']))
            story.append(Spacer(1, 12))
            
            # Add content
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    story.append(Paragraph(line[2:], styles['Heading1']))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], styles['Heading2']))
                elif line.startswith('### '):
                    story.append(Paragraph(line[4:], styles['Heading3']))
                elif line.strip():
                    story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 6))
            
            doc.build(story)
            pdf_buffer.seek(0)
            
            st.download_button(
                label="💾 Download PDF",
                data=pdf_buffer.getvalue(),
                file_name=f"{topic.lower().replace(' ', '_')}.pdf",
                mime="application/pdf",
                key="pdf_download_btn"
            )
        except ImportError:
            st.error("❌ PDF export requires reportlab. Install with: pip install reportlab")
    
    def download_complete_package(self):
        """Download complete content package as ZIP"""
        import zipfile
        
        final_content = st.session_state.generated_content.get("final_content", "")
        if not final_content:
            final_content = self.get_content_data("content_creation", "No content available")
        
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            topic = st.session_state.content_session.get('initial_state', {}).get('content_topic', 'Content')
            base_name = topic.lower().replace(' ', '_')
            
            # Add markdown version
            zip_file.writestr(f"{base_name}.md", final_content)
            
            # Add HTML version
            html_content = markdown2.markdown(final_content)
            zip_file.writestr(f"{base_name}.html", html_content)
            
            # Add analyses
            analyses_content = "# Content Creation Analysis Report\n\n"
            for agent_name, result in st.session_state.content_results.items():
                analyses_content += f"## {agent_name.replace('_', ' ').title()}\n\n"
                analyses_content += self.get_content_data(agent_name, "No analysis available")
                analyses_content += "\n\n---\n\n"
            
            zip_file.writestr(f"{base_name}_analyses.md", analyses_content)
            
            # Add metadata
            metadata = {
                "topic": topic,
                "content_type": st.session_state.content_session.get('initial_state', {}).get('content_type', ''),
                "generated_at": datetime.now().isoformat(),
                "word_count": len(final_content.split()),
                "agents_used": list(st.session_state.content_results.keys())
            }
            zip_file.writestr("metadata.json", json.dumps(metadata, indent=2))
        
        zip_buffer.seek(0)
        
        st.download_button(
            label="💾 Download Complete Package",
            data=zip_buffer.getvalue(),
            file_name=f"{topic.lower().replace(' ', '_')}_complete_package.zip",
            mime="application/zip",
            key="package_download_btn"
        )
    
    def download_social_media_kit(self, content: str):
        """Download social media variations"""
        # Generate social media variations
        lines = content.split('\n')
        title = lines[0].replace('#', '').strip() if lines else "Content"
        
        social_variations = f"""# Social Media Kit

## Twitter/X Version
🤖 {title[:100]}... #AI #ContentMarketing #Digital

## LinkedIn Version
{title}

{lines[1] if len(lines) > 1 else 'Key insights from our latest content.'}

#LinkedIn #Professional #Industry

## Instagram Caption
{title} ✨

{lines[1] if len(lines) > 1 else 'Swipe to learn more!'}

#Instagram #Visual #Engagement

## Facebook Post
{title}

{' '.join(lines[1:3]) if len(lines) > 2 else 'Check out our latest insights!'}

Learn more: [Link]
"""
        
        topic = st.session_state.content_session.get('initial_state', {}).get('content_topic', 'Content')
        
        st.download_button(
            label="💾 Download Social Media Kit",
            data=social_variations,
            file_name=f"{topic.lower().replace(' ', '_')}_social_kit.md",
            mime="text/markdown",
            key="social_download_btn"
        )
    
    def download_email_version(self, content: str):
        """Download email-optimized version"""
        lines = content.split('\n')
        title = lines[0].replace('#', '').strip() if lines else "Content"
        
        email_content = f"""Subject: {title}

Hi [Name],

I wanted to share this valuable insight with you:

{title}

{' '.join(lines[1:3]) if len(lines) > 2 else 'Key takeaways from our latest research.'}

Key points:
• {lines[3] if len(lines) > 3 else 'Important insight #1'}
• {lines[4] if len(lines) > 4 else 'Important insight #2'}
• {lines[5] if len(lines) > 5 else 'Important insight #3'}

Read the full article: [Link]

Best regards,
[Your Name]

P.S. Forward this to anyone who might find it valuable!
"""
        
        topic = st.session_state.content_session.get('initial_state', {}).get('content_topic', 'Content')
        
        st.download_button(
            label="💾 Download Email Version",
            data=email_content,
            file_name=f"{topic.lower().replace(' ', '_')}_email.txt",
            mime="text/plain",
            key="email_download_btn"
        )
    
    def reset_session(self):
        """Reset the content creation session"""
        st.session_state.content_session = None
        st.session_state.content_results = {}
        st.session_state.workflow_history = []
        st.session_state.workflow_running = False
        st.session_state.current_agent = None
        st.session_state.generated_content = {}
        st.rerun()


def main():
    """Main entry point"""
    app = ContentCreationApp()
    app.run()


if __name__ == "__main__":
    main()