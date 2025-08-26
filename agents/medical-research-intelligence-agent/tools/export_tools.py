"""
Medical Report Export Tools for MARIA
Healthcare-specific report generation and export utilities
"""

import os
import json
import csv
from io import StringIO, BytesIO
from typing import Dict, List, Any, Optional
from datetime import datetime
import tempfile

try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class MedicalReportExporter:
    """Export medical research reports in various formats"""
    
    def __init__(self, research_data: Dict[str, Any]):
        """
        Initialize medical report exporter
        
        Args:
            research_data: Medical research data to export
        """
        self.research_data = research_data
        self.timestamp = datetime.now()
        self.medical_disclaimer = """
IMPORTANT MEDICAL DISCLAIMER:
This report contains AI-generated medical research content for informational and research purposes only. 
All information requires validation by qualified healthcare professionals before any clinical application.
This content does not constitute medical advice and should not be used for patient care decisions without 
proper medical oversight and validation.
"""
    
    def export_clinical_pdf_report(self) -> bytes:
        """
        Export clinical PDF report
        
        Returns:
            PDF report as bytes
        """
        if not PDF_AVAILABLE:
            raise ImportError("fpdf2 not available. Install with: pip install fpdf2")
        
        try:
            pdf = FPDF()
            pdf.add_page()
            
            # Header
            pdf.set_font('Arial', 'B', 16)
            pdf.cell(0, 10, 'MARIA - Medical Research Intelligence Report', ln=True, align='C')
            
            pdf.set_font('Arial', '', 10)
            pdf.cell(0, 8, f'Generated: {self.timestamp.strftime("%Y-%m-%d %H:%M:%S")}', ln=True, align='C')
            pdf.ln(5)
            
            # Medical disclaimer
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(0, 8, 'MEDICAL DISCLAIMER', ln=True)
            pdf.set_font('Arial', '', 9)
            disclaimer_lines = self.medical_disclaimer.strip().split('\n')
            for line in disclaimer_lines:
                if line.strip():
                    pdf.cell(0, 5, line.strip(), ln=True)
            pdf.ln(5)
            
            # Research context
            research_state = self.research_data.get('research_state', {})
            if research_state:
                pdf.set_font('Arial', 'B', 12)
                pdf.cell(0, 8, 'Research Context', ln=True)
                pdf.set_font('Arial', '', 10)
                
                context_items = [
                    ('Researcher ID:', research_state.get('researcher_id', 'Not specified')),
                    ('Disease Focus:', research_state.get('disease_focus', 'Not specified')),
                    ('Research Type:', research_state.get('research_type', 'Not specified')),
                    ('Target Population:', research_state.get('target_population', 'Not specified')),
                    ('Research Depth:', research_state.get('research_depth', 'Not specified'))
                ]
                
                for label, value in context_items:
                    pdf.cell(50, 6, label, 0, 0)
                    pdf.cell(0, 6, str(value), ln=True)
                pdf.ln(5)
            
            # Conversation content
            conversation = self.research_data.get('conversation', [])
            if conversation:
                pdf.set_font('Arial', 'B', 12)
                pdf.cell(0, 8, 'Medical Research Findings', ln=True)
                
                for i, message in enumerate(conversation):
                    if message.get('sender') == 'assistant':
                        pdf.set_font('Arial', 'B', 10)
                        pdf.cell(0, 6, f'Medical Analysis {i+1}:', ln=True)
                        
                        pdf.set_font('Arial', '', 9)
                        content = message.get('content', '')
                        
                        # Split content into lines that fit
                        words = content.split()
                        current_line = ""
                        for word in words:
                            if len(current_line + " " + word) < 85:  # Approximate character limit
                                current_line += " " + word if current_line else word
                            else:
                                if current_line:
                                    pdf.cell(0, 4, current_line, ln=True)
                                current_line = word
                        
                        if current_line:
                            pdf.cell(0, 4, current_line, ln=True)
                        
                        # Add confidence score if available
                        confidence = message.get('confidence_score')
                        if confidence:
                            pdf.set_font('Arial', 'I', 8)
                            pdf.cell(0, 4, f'Confidence Score: {confidence:.2f}', ln=True)
                        
                        pdf.ln(3)
            
            # Footer
            pdf.ln(10)
            pdf.set_font('Arial', 'I', 8)
            pdf.cell(0, 4, 'This report was generated by MARIA (Medical Research Intelligence Agent)', ln=True, align='C')
            pdf.cell(0, 4, 'All medical content requires professional validation', ln=True, align='C')
            
            return pdf.output(dest='S').encode('latin-1')
        
        except Exception as e:
            raise Exception(f"Error creating clinical PDF report: {str(e)}")
    
    def export_research_csv(self) -> str:
        """
        Export research data as CSV
        
        Returns:
            CSV data as string
        """
        try:
            output = StringIO()
            writer = csv.writer(output)
            
            # Header
            writer.writerow(['Medical Research Data Export'])
            writer.writerow(['Generated:', self.timestamp.strftime("%Y-%m-%d %H:%M:%S")])
            writer.writerow([])
            
            # Research context
            research_state = self.research_data.get('research_state', {})
            if research_state:
                writer.writerow(['Research Context'])
                writer.writerow(['Field', 'Value'])
                
                for key, value in research_state.items():
                    if isinstance(value, (str, int, float)):
                        writer.writerow([key.replace('_', ' ').title(), str(value)])
                
                writer.writerow([])
            
            # Conversation data
            conversation = self.research_data.get('conversation', [])
            if conversation:
                writer.writerow(['Medical Research Conversation'])
                writer.writerow(['Timestamp', 'Sender', 'Content', 'Confidence Score', 'Medical Context'])
                
                for message in conversation:
                    timestamp = message.get('timestamp', '')
                    sender = message.get('sender', '')
                    content = message.get('content', '').replace('\n', ' | ')  # Replace newlines
                    confidence = message.get('confidence_score', '')
                    medical_context = str(message.get('medical_context', {}))
                    
                    writer.writerow([timestamp, sender, content, confidence, medical_context])
            
            # Medical disclaimer
            writer.writerow([])
            writer.writerow(['Medical Disclaimer'])
            for line in self.medical_disclaimer.strip().split('\n'):
                if line.strip():
                    writer.writerow([line.strip()])
            
            return output.getvalue()
        
        except Exception as e:
            raise Exception(f"Error creating research CSV: {str(e)}")
    
    def export_literature_review(self) -> str:
        """
        Export literature review in markdown format
        
        Returns:
            Literature review as markdown string
        """
        try:
            research_state = self.research_data.get('research_state', {})
            conversation = self.research_data.get('conversation', [])
            
            md_content = f"""# Medical Literature Review
## {research_state.get('disease_focus', 'Medical Research')} Analysis

**Generated by MARIA (Medical Research Intelligence Agent)**  
**Date:** {self.timestamp.strftime("%Y-%m-%d %H:%M:%S")}

---

## Medical Disclaimer
{self.medical_disclaimer}

---

## Research Overview

### Research Context
- **Researcher ID:** {research_state.get('researcher_id', 'Not specified')}
- **Disease Focus:** {research_state.get('disease_focus', 'Not specified')}
- **Research Type:** {research_state.get('research_type', 'Not specified')}
- **Target Population:** {research_state.get('target_population', 'Not specified')}
- **Research Depth:** {research_state.get('research_depth', 'Not specified')}

### Research Methodology
This literature review was conducted using AI-assisted medical research tools, including:
- Medical literature database analysis
- Clinical evidence synthesis
- Treatment efficacy evaluation
- Safety profile assessment

---

## Medical Research Findings

"""
            
            # Add conversation content
            for i, message in enumerate(conversation):
                if message.get('sender') == 'assistant':
                    content = message.get('content', '')
                    confidence = message.get('confidence_score', 0)
                    
                    md_content += f"""### Research Finding {i+1}

{content}

**Confidence Score:** {confidence:.2f}/1.0

---

"""
            
            # Add validation notes
            md_content += f"""## Validation and Quality Assessment

### Evidence Quality
- **AI Analysis Confidence:** Variable (see individual sections)
- **Data Sources:** Medical literature databases, clinical guidelines
- **Validation Status:** Requires professional medical review
- **Last Updated:** {self.timestamp.strftime("%Y-%m-%d")}

### Limitations
- AI-generated content requires human validation
- Limited to available literature in training data
- May not reflect most recent medical developments
- Requires clinical context for practical application

### Recommendations for Use
1. **Professional Review:** All findings should be reviewed by qualified healthcare professionals
2. **Clinical Validation:** Content should be validated against current medical standards
3. **Patient Safety:** Do not use for direct patient care without medical oversight
4. **Continuous Updates:** Regular updates needed as medical knowledge evolves

---

## References and Sources

### Primary Sources
- Medical literature databases (PubMed, Cochrane, EMBASE)
- Professional medical guidelines
- Clinical trial registries
- Peer-reviewed medical journals

### AI Analysis Tools
- MARIA (Medical Research Intelligence Agent)
- AutoGen medical conversation framework
- Healthcare-specific language models

---

**Report Generated by:** MARIA - Medical Research Intelligence Agent  
**Institution:** MediSyn Labs  
**Contact:** For questions about this analysis, consult with qualified medical professionals  
**Version:** {datetime.now().strftime("%Y.%m.%d")}
"""
            
            return md_content
        
        except Exception as e:
            raise Exception(f"Error creating literature review: {str(e)}")
    
    def export_prisma_report(self) -> str:
        """
        Export PRISMA-style systematic review report
        
        Returns:
            PRISMA report as markdown string
        """
        try:
            research_state = self.research_data.get('research_state', {})
            
            prisma_content = f"""# PRISMA-Style Systematic Review Report
## {research_state.get('disease_focus', 'Medical Condition')} Research Analysis

**Review Protocol:** AI-Assisted Medical Literature Analysis  
**Review Date:** {self.timestamp.strftime("%Y-%m-%d")}  
**Generated by:** MARIA (Medical Research Intelligence Agent)

---

## Medical Disclaimer
{self.medical_disclaimer}

---

## Abstract

### Background
This systematic review examines current evidence for {research_state.get('disease_focus', 'the specified medical condition')} using AI-assisted literature analysis.

### Objectives
To synthesize current medical evidence regarding {research_state.get('research_type', 'treatment and management')} for {research_state.get('disease_focus', 'the target condition')} in {research_state.get('target_population', 'the specified population')}.

### Methods
AI-assisted systematic review using medical literature databases and evidence synthesis algorithms.

### Results
Evidence synthesis provided with confidence scoring and quality assessment.

### Conclusions
All findings require validation by qualified medical professionals before clinical application.

---

## 1. Introduction

### Research Question
What is the current evidence for {research_state.get('research_type', 'treatment approaches')} in {research_state.get('disease_focus', 'the target condition')}?

### Objectives
- Synthesize current medical literature
- Assess treatment efficacy and safety
- Identify research gaps and limitations
- Provide evidence-based recommendations

---

## 2. Methods

### Search Strategy
- **Databases:** Medical literature databases (AI training data)
- **Search Terms:** Condition-specific medical terminology
- **Time Period:** Current medical knowledge base
- **Language:** English

### Inclusion Criteria
- Peer-reviewed medical literature
- Clinical studies and trials
- Professional medical guidelines
- Evidence-based medical content

### Exclusion Criteria
- Non-medical sources
- Unverified medical claims
- Outdated medical practices
- Non-evidence-based content

### Data Extraction
- Study characteristics and design
- Patient populations and interventions
- Clinical outcomes and safety data
- Evidence quality assessment

---

## 3. Results

### Study Selection
AI-assisted analysis of available medical literature in training database.

### Study Characteristics
Studies included various research designs relevant to {research_state.get('disease_focus', 'the condition')}.

### Evidence Synthesis
"""
            
            # Add conversation findings
            conversation = self.research_data.get('conversation', [])
            for i, message in enumerate(conversation):
                if message.get('sender') == 'assistant':
                    content = message.get('content', '')
                    confidence = message.get('confidence_score', 0)
                    
                    prisma_content += f"""

#### Finding {i+1}
{content}

**Evidence Level:** {self._determine_evidence_level(confidence)}  
**Confidence Score:** {confidence:.2f}/1.0
"""
            
            prisma_content += f"""

---

## 4. Discussion

### Summary of Evidence
The AI-assisted analysis provides preliminary evidence synthesis for {research_state.get('disease_focus', 'the medical condition')}. All findings require professional medical validation.

### Limitations
- AI-generated content requires human validation
- Limited to training data knowledge base
- May not reflect most recent developments
- Requires clinical context for application

### Clinical Implications
All clinical implications must be validated by qualified healthcare professionals before implementation.

---

## 5. Conclusions

### Main Findings
Evidence synthesis completed using AI assistance with {research_state.get('research_depth', 'intermediate')} depth analysis.

### Implications for Practice
All practice implications require validation by medical professionals.

### Implications for Research
Further human-led systematic reviews recommended for clinical validation.

---

## 6. Funding and Conflicts of Interest

**Funding:** AI-assisted research tool development  
**Conflicts of Interest:** None declared  
**Disclosure:** This analysis was generated using artificial intelligence and requires medical professional validation.

---

## 7. References

### Primary Literature Sources
- Medical literature databases (PubMed, Cochrane, EMBASE)
- Professional medical guidelines
- Clinical trial registries
- Peer-reviewed medical journals

### AI Analysis Framework
- MARIA (Medical Research Intelligence Agent)
- AutoGen healthcare conversation system
- Medical language model analysis

---

**PRISMA Compliance Note:** This report follows PRISMA-style formatting but is AI-generated and requires validation by qualified systematic review experts and medical professionals.

**Report Metadata:**
- Generator: MARIA v{datetime.now().strftime("%Y.%m")}
- Analysis Date: {self.timestamp.strftime("%Y-%m-%d %H:%M:%S")}
- Validation Status: Pending professional review
"""
            
            return prisma_content
        
        except Exception as e:
            raise Exception(f"Error creating PRISMA report: {str(e)}")
    
    def _determine_evidence_level(self, confidence_score: float) -> str:
        """Determine evidence level based on confidence score"""
        if confidence_score >= 0.9:
            return "High"
        elif confidence_score >= 0.7:
            return "Moderate"
        elif confidence_score >= 0.5:
            return "Low"
        else:
            return "Very Low"
    
    def export_word_document(self) -> bytes:
        """
        Export medical report as Word document
        
        Returns:
            Word document as bytes
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")
        
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('MARIA - Medical Research Intelligence Report', 0)
            
            # Subtitle
            research_state = self.research_data.get('research_state', {})
            subtitle = f"Medical Analysis: {research_state.get('disease_focus', 'Healthcare Research')}"
            doc.add_heading(subtitle, level=1)
            
            # Metadata
            doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Researcher: {research_state.get('researcher_id', 'Not specified')}")
            
            # Medical disclaimer
            disclaimer_para = doc.add_paragraph()
            disclaimer_para.add_run('MEDICAL DISCLAIMER: ').bold = True
            disclaimer_para.add_run(self.medical_disclaimer.strip())
            
            # Research context
            doc.add_heading('Research Context', level=2)
            context_table = doc.add_table(rows=1, cols=2)
            context_table.style = 'Table Grid'
            
            hdr_cells = context_table.rows[0].cells
            hdr_cells[0].text = 'Field'
            hdr_cells[1].text = 'Value'
            
            context_items = [
                ('Disease Focus', research_state.get('disease_focus', 'Not specified')),
                ('Research Type', research_state.get('research_type', 'Not specified')),
                ('Target Population', research_state.get('target_population', 'Not specified')),
                ('Research Depth', research_state.get('research_depth', 'Not specified'))
            ]
            
            for field, value in context_items:
                row_cells = context_table.add_row().cells
                row_cells[0].text = field
                row_cells[1].text = str(value)
            
            # Medical findings
            doc.add_heading('Medical Research Findings', level=2)
            
            conversation = self.research_data.get('conversation', [])
            for i, message in enumerate(conversation):
                if message.get('sender') == 'assistant':
                    doc.add_heading(f'Medical Analysis {i+1}', level=3)
                    
                    content = message.get('content', '')
                    doc.add_paragraph(content)
                    
                    confidence = message.get('confidence_score')
                    if confidence:
                        confidence_para = doc.add_paragraph()
                        confidence_para.add_run('Confidence Score: ').italic = True
                        confidence_para.add_run(f'{confidence:.2f}/1.0')
            
            # Save to bytes
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer.getvalue()
        
        except Exception as e:
            raise Exception(f"Error creating Word document: {str(e)}")


def create_medical_report_exporter(research_data: Dict[str, Any]) -> MedicalReportExporter:
    """
    Create medical report exporter
    
    Args:
        research_data: Medical research data
        
    Returns:
        MedicalReportExporter instance
    """
    return MedicalReportExporter(research_data)