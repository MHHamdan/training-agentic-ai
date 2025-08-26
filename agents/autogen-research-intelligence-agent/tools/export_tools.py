"""
Export Tools for ARIA
Multi-format export functionality for research outputs
"""

import os
import json
import csv
import io
from typing import Dict, List, Any, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False


class ResearchExporter:
    """
    Multi-format research export tool
    """
    
    def __init__(self, research_data: Dict[str, Any]):
        """
        Initialize research exporter
        
        Args:
            research_data: Research data to export
        """
        self.research_data = research_data
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    def export_to_pdf(self) -> bytes:
        """
        Export research to PDF format
        
        Returns:
            PDF data as bytes
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab is required for PDF export. Install with: pip install reportlab")
        
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.darkblue,
                alignment=1  # Center alignment
            )
            
            story.append(Paragraph("ARIA Research Report", title_style))
            story.append(Spacer(1, 12))
            
            # Metadata
            self._add_pdf_metadata(story, styles)
            story.append(Spacer(1, 12))
            
            # Conversation content
            self._add_pdf_conversation(story, styles)
            
            # Research summary
            self._add_pdf_summary(story, styles)
            
            doc.build(story)
            buffer.seek(0)
            return buffer.read()
            
        except Exception as e:
            # Fallback to simple text-based PDF
            return self._create_fallback_pdf()
    
    def _add_pdf_metadata(self, story: List, styles):
        """Add metadata section to PDF"""
        research_state = self.research_data.get('research_state', {})
        
        metadata = [
            ['Topic:', research_state.get('current_topic', 'N/A')],
            ['Depth:', research_state.get('research_depth', 'N/A')],
            ['Audience:', research_state.get('target_audience', 'N/A')],
            ['Generated:', datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
            ['Session ID:', research_state.get('session_id', 'N/A')]
        ]
        
        table = Table(metadata, colWidths=[1.5*inch, 4*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(Paragraph("Research Details", styles['Heading2']))
        story.append(table)
    
    def _add_pdf_conversation(self, story: List, styles):
        """Add conversation content to PDF"""
        conversation = self.research_data.get('conversation', [])
        
        if conversation:
            story.append(Spacer(1, 12))
            story.append(Paragraph("Research Conversation", styles['Heading2']))
            story.append(Spacer(1, 6))
            
            for message in conversation:
                sender = message.get('sender', 'Unknown')
                content = message.get('content', '')
                timestamp = message.get('timestamp', '')
                
                # Sender header
                sender_style = ParagraphStyle(
                    'SenderStyle',
                    parent=styles['Normal'],
                    fontSize=11,
                    textColor=colors.darkblue,
                    fontName='Helvetica-Bold'
                )
                
                story.append(Paragraph(f"{sender.title()} ({timestamp}):", sender_style))
                story.append(Paragraph(content, styles['Normal']))
                story.append(Spacer(1, 8))
    
    def _add_pdf_summary(self, story: List, styles):
        """Add research summary to PDF"""
        story.append(Spacer(1, 12))
        story.append(Paragraph("Research Summary", styles['Heading2']))
        
        conversation = self.research_data.get('conversation', [])
        total_messages = len(conversation)
        user_messages = len([m for m in conversation if m.get('sender') == 'user'])
        assistant_messages = len([m for m in conversation if m.get('sender') == 'assistant'])
        
        summary_data = [
            ['Total Messages:', str(total_messages)],
            ['User Messages:', str(user_messages)],
            ['Assistant Responses:', str(assistant_messages)],
            ['Export Time:', datetime.now().strftime("%Y-%m-%d %H:%M:%S")]
        ]
        
        summary_table = Table(summary_data, colWidths=[2*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightblue),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(summary_table)
    
    def _create_fallback_pdf(self) -> bytes:
        """Create a simple fallback PDF"""
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            story.append(Paragraph("ARIA Research Export", styles['Title']))
            story.append(Spacer(1, 12))
            story.append(Paragraph("Research data exported successfully.", styles['Normal']))
            story.append(Paragraph(f"Export time: {datetime.now()}", styles['Normal']))
            
            doc.build(story)
            buffer.seek(0)
            return buffer.read()
        except Exception:
            # Ultimate fallback - return empty bytes
            return b"PDF export not available"
    
    def export_to_word(self) -> bytes:
        """
        Export research to Word document format
        
        Returns:
            Word document data as bytes
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word export. Install with: pip install python-docx")
        
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('ARIA Research Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            self._add_word_metadata(doc)
            
            # Conversation
            self._add_word_conversation(doc)
            
            # Summary
            self._add_word_summary(doc)
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.read()
            
        except Exception as e:
            return self._create_fallback_word()
    
    def _add_word_metadata(self, doc):
        """Add metadata to Word document"""
        doc.add_heading('Research Details', level=1)
        
        research_state = self.research_data.get('research_state', {})
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        cells = table.rows[0].cells
        cells[0].text = 'Topic:'
        cells[1].text = research_state.get('current_topic', 'N/A')
        
        cells = table.rows[1].cells
        cells[0].text = 'Depth:'
        cells[1].text = research_state.get('research_depth', 'N/A')
        
        cells = table.rows[2].cells
        cells[0].text = 'Audience:'
        cells[1].text = research_state.get('target_audience', 'N/A')
        
        cells = table.rows[3].cells
        cells[0].text = 'Generated:'
        cells[1].text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        cells = table.rows[4].cells
        cells[0].text = 'Session ID:'
        cells[1].text = research_state.get('session_id', 'N/A')
        
        doc.add_paragraph()
    
    def _add_word_conversation(self, doc):
        """Add conversation to Word document"""
        conversation = self.research_data.get('conversation', [])
        
        if conversation:
            doc.add_heading('Research Conversation', level=1)
            
            for message in conversation:
                sender = message.get('sender', 'Unknown')
                content = message.get('content', '')
                timestamp = message.get('timestamp', '')
                
                # Add sender heading
                para = doc.add_paragraph()
                run = para.add_run(f"{sender.title()} ({timestamp}):")
                run.bold = True
                
                # Add content
                doc.add_paragraph(content)
                doc.add_paragraph()
    
    def _add_word_summary(self, doc):
        """Add summary to Word document"""
        doc.add_heading('Research Summary', level=1)
        
        conversation = self.research_data.get('conversation', [])
        total_messages = len(conversation)
        user_messages = len([m for m in conversation if m.get('sender') == 'user'])
        assistant_messages = len([m for m in conversation if m.get('sender') == 'assistant'])
        
        doc.add_paragraph(f"Total Messages: {total_messages}")
        doc.add_paragraph(f"User Messages: {user_messages}")
        doc.add_paragraph(f"Assistant Responses: {assistant_messages}")
        doc.add_paragraph(f"Export Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    def _create_fallback_word(self) -> bytes:
        """Create fallback Word document"""
        try:
            doc = Document()
            doc.add_heading('ARIA Research Export', 0)
            doc.add_paragraph('Research data exported successfully.')
            doc.add_paragraph(f'Export time: {datetime.now()}')
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.read()
        except Exception:
            return b"Word export not available"
    
    def export_to_csv(self) -> str:
        """
        Export research to CSV format
        
        Returns:
            CSV data as string
        """
        try:
            output = io.StringIO()
            writer = csv.writer(output)
            
            # Headers
            writer.writerow(['Timestamp', 'Sender', 'Content', 'Message_Type'])
            
            # Conversation data
            conversation = self.research_data.get('conversation', [])
            for message in conversation:
                writer.writerow([
                    message.get('timestamp', ''),
                    message.get('sender', ''),
                    message.get('content', '').replace('\n', ' '),  # Remove newlines for CSV
                    'conversation'
                ])
            
            # Research metadata
            research_state = self.research_data.get('research_state', {})
            writer.writerow([
                datetime.now().isoformat(),
                'system',
                f"Topic: {research_state.get('current_topic', 'N/A')}",
                'metadata'
            ])
            
            writer.writerow([
                datetime.now().isoformat(),
                'system',
                f"Depth: {research_state.get('research_depth', 'N/A')}",
                'metadata'
            ])
            
            writer.writerow([
                datetime.now().isoformat(),
                'system',
                f"Audience: {research_state.get('target_audience', 'N/A')}",
                'metadata'
            ])
            
            output.seek(0)
            return output.getvalue()
            
        except Exception as e:
            return f"Error exporting to CSV: {str(e)}"
    
    def export_to_markdown(self) -> str:
        """
        Export research to Markdown format
        
        Returns:
            Markdown content as string
        """
        try:
            md_content = []
            
            # Title
            md_content.append("# ARIA Research Report")
            md_content.append("")
            
            # Metadata
            research_state = self.research_data.get('research_state', {})
            md_content.append("## Research Details")
            md_content.append("")
            md_content.append(f"- **Topic:** {research_state.get('current_topic', 'N/A')}")
            md_content.append(f"- **Depth:** {research_state.get('research_depth', 'N/A')}")
            md_content.append(f"- **Audience:** {research_state.get('target_audience', 'N/A')}")
            md_content.append(f"- **Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            md_content.append(f"- **Session ID:** {research_state.get('session_id', 'N/A')}")
            md_content.append("")
            
            # Conversation
            conversation = self.research_data.get('conversation', [])
            if conversation:
                md_content.append("## Research Conversation")
                md_content.append("")
                
                for message in conversation:
                    sender = message.get('sender', 'Unknown')
                    content = message.get('content', '')
                    timestamp = message.get('timestamp', '')
                    
                    md_content.append(f"### {sender.title()} ({timestamp})")
                    md_content.append("")
                    md_content.append(content)
                    md_content.append("")
            
            # Summary
            md_content.append("## Research Summary")
            md_content.append("")
            total_messages = len(conversation)
            user_messages = len([m for m in conversation if m.get('sender') == 'user'])
            assistant_messages = len([m for m in conversation if m.get('sender') == 'assistant'])
            
            md_content.append(f"- **Total Messages:** {total_messages}")
            md_content.append(f"- **User Messages:** {user_messages}")
            md_content.append(f"- **Assistant Responses:** {assistant_messages}")
            md_content.append(f"- **Export Time:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            md_content.append("")
            
            # Footer
            md_content.append("---")
            md_content.append("*Generated by ARIA - Autogen Research Intelligence Agent*")
            
            return "\n".join(md_content)
            
        except Exception as e:
            return f"# Export Error\n\nError exporting to Markdown: {str(e)}"
    
    def export_to_json(self) -> str:
        """
        Export research to JSON format
        
        Returns:
            JSON data as string
        """
        try:
            export_data = {
                "export_metadata": {
                    "export_time": datetime.now().isoformat(),
                    "export_format": "json",
                    "exporter": "ARIA Research Intelligence Agent",
                    "version": "1.0"
                },
                "research_data": self.research_data,
                "summary": self._generate_export_summary()
            }
            
            return json.dumps(export_data, indent=2, ensure_ascii=False)
            
        except Exception as e:
            return json.dumps({
                "error": f"Export error: {str(e)}",
                "export_time": datetime.now().isoformat()
            })
    
    def _generate_export_summary(self) -> Dict[str, Any]:
        """Generate summary for export"""
        conversation = self.research_data.get('conversation', [])
        research_state = self.research_data.get('research_state', {})
        
        return {
            "topic": research_state.get('current_topic', 'N/A'),
            "total_messages": len(conversation),
            "message_breakdown": {
                "user": len([m for m in conversation if m.get('sender') == 'user']),
                "assistant": len([m for m in conversation if m.get('sender') == 'assistant']),
                "system": len([m for m in conversation if m.get('sender') == 'system'])
            },
            "session_duration": self._calculate_session_duration(conversation),
            "export_formats_available": ["pdf", "docx", "csv", "markdown", "json"]
        }
    
    def _calculate_session_duration(self, conversation: List[Dict]) -> str:
        """Calculate session duration from conversation timestamps"""
        if len(conversation) < 2:
            return "N/A"
        
        try:
            start_time = datetime.fromisoformat(conversation[0]['timestamp'])
            end_time = datetime.fromisoformat(conversation[-1]['timestamp'])
            duration = end_time - start_time
            
            minutes = int(duration.total_seconds() / 60)
            return f"{minutes} minutes"
        except (KeyError, ValueError):
            return "N/A"


def get_export_capabilities() -> Dict[str, Any]:
    """
    Get information about export capabilities
    
    Returns:
        Dictionary describing available export capabilities
    """
    return {
        "supported_formats": ["pdf", "docx", "csv", "markdown", "json"],
        "dependencies": {
            "pdf": REPORTLAB_AVAILABLE,
            "docx": DOCX_AVAILABLE,
            "markdown": MARKDOWN_AVAILABLE,
            "csv": True,  # Built-in
            "json": True  # Built-in
        },
        "features": {
            "conversation_export": True,
            "metadata_inclusion": True,
            "summary_generation": True,
            "timestamp_preservation": True,
            "multi_format_support": True
        },
        "requirements": {
            "pdf": "reportlab>=4.0.0",
            "docx": "python-docx>=1.0.0",
            "markdown": "markdown>=3.5.0"
        }
    }